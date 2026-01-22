from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import configparser

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)

config = configparser.ConfigParser()
config.read("config.properties")
# -------------------------------------------------
# 1️⃣ DATA (DICT DRIVES EVERYTHING)
# -------------------------------------------------
# table_data = {
#     "ticket_data": {
#         "ticket": ["INC-10234", "INC-10235", "INC-10236"],
#         "summary": [
#             "Login page error: unable to load the login page due to missing session token, which is critical for user authentication",
#             "API timeout issue while calling the external payment gateway, resulting in failed transactions for multiple users",
#             "Database connection failed intermittently due to high load on primary DB server, needs urgent investigation"
#         ]
#     },
#     "ticket_data2": {
#         "ticket": ["INC-20234", "INC-20235", "INC-20236"],
#         "summary": [
#             "Cache invalidation issue causing stale data",
#             "File upload fails for large attachments",
#             "Unexpected logout during session refresh"
#         ]
#     },
#     "ticket_data3": {
#         "ticket": ["INC-30234", "INC-30235", "INC-30236"],
#         "summary": [
#             "UI alignment breaks on mobile view",
#             "Permission mismatch for admin users",
#             "Background job retry logic missing"
#         ]
#     }
# }

table_data = dict(
    map(
        lambda table: (
            table,
            dict(
                map(
                    lambda col: (col.strip(), ""),  # column name -> empty string
                    config.get("table_placeholders", table).split(",")
                )
            )
        ),
        config["table_placeholders"]
    )
)

placeholder_data = dict(
    map(
        lambda np: ("{{" + np.strip().upper() + "}}", ''),
        config.get("paragraph_placeholders", "fields").split(",")
    )
)

def add_placeholder_data():
    for key, value in placeholder_data.items():
        placeholder_data[key] = input(f"Enter value for {key}: ")
def add_table_data():
    for table_name, columns in table_data.items():
        print(f"\n--- Enter data for table '{table_name}' ---")
        num_rows = int(input("How many rows? "))
        for col_name in columns:
            table_data[table_name][col_name] = []
        for row_num in range(1, num_rows + 1):
            print(f"  Row {row_num}:")
            for col_name in columns:
                val = input(f"    {col_name}: ")
                table_data[table_name][col_name].append(val)

# -------------------------------------------------
# 2️⃣ CONSTANTS
# -------------------------------------------------

PAGE_WIDTH = config.getfloat("table", "page_width")
MIN_COL_WIDTH = config.getfloat("table", "min_col_width")
CHAR_PER_INCH = config.getint("table", "char_per_inch")

# -------------------------------------------------
# 3️⃣ HELPERS
# -------------------------------------------------
def calculate_column_widths(df):
    num_cols = len(df.columns)
    max_col_width = PAGE_WIDTH * 0.6  # No single column takes more than 60%
    
    col_lengths = {}
    for col in df.columns:
        max_len = max(df[col].astype(str).map(len).max(), len(col))
        col_lengths[col] = min(max_len, 50)  # Cap character length at 50

    raw_widths = {
        col: max(MIN_COL_WIDTH, min(length / CHAR_PER_INCH, max_col_width))
        for col, length in col_lengths.items()
    }

    total_width = sum(raw_widths.values())
    if total_width > PAGE_WIDTH:
        scale = PAGE_WIDTH / total_width
        return {col: max(MIN_COL_WIDTH, width * scale) for col, width in raw_widths.items()}
    return raw_widths


def set_column_widths(table, widths):
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Inches(width)
        table.columns[i].width = Inches(width)

def insert_table_at_placeholder(doc, placeholder, table_content):
    df = pd.DataFrame(table_content)

    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = ""  # remove placeholder

            table = doc.add_table(rows=1, cols=len(df.columns))

            # Header
            for i, col in enumerate(df.columns):
                cell = table.cell(0, i)
                cell.text = col.upper()
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Rows
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, col in enumerate(df.columns):
                    row_cells[i].text = str(row[col])
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Apply borders after all rows added
            set_table_borders(table)

            # Insert table at placeholder position
            para._p.addnext(table._tbl)

def replace_in_runs(paragraph, key, value):
    for run in paragraph.runs:
        if key in run.text:
            run.text = run.text.replace(key, value)

def update_docx(doc,values):
    for paragraph in doc.paragraphs:
        for key, value in values.items():
            if key in paragraph.text:
                replace_in_runs(paragraph, key, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in values.items():
                    if key in cell.text:
                        for para in cell.paragraphs:
                            if key in para.text:
                                replace_in_runs(para, key, value)
# -------------------------------------------------
# 4️⃣ MAIN
# -------------------------------------------------
def main():
    doc = Document(config["paths"]["template"])

    add_placeholder_data()
    add_table_data()
    update_docx(doc, placeholder_data)
    for key, value in table_data.items():
        placeholder = f"{{{{{key.upper()}}}}}"  # {{RELEASE_TABLE}}
        insert_table_at_placeholder(doc, placeholder, value)

    doc.save(f"Release_Note_Release_{placeholder_data["{{SNO}}"]}_.docx")
    print("✅ Document generated successfully")


if __name__ == "__main__":
    main()
